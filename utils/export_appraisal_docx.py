"""
Export quality appraisal results to a formatted Word document (.docx).

Produces a document styled to match the REST Quality Assessment Tool format,
with one section per paper containing the 20-criterion rating table,
quality score, and strengths/limitations narrative.
"""

from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.appraisal_schemas import PaperAppraisal


# ── Colour palette ───────────────────────────────────────────────────────────

COLOUR_HEADER_BG   = RGBColor(0x1B, 0x2A, 0x4A)   # Deep navy — REST brand
COLOUR_HEADER_FG   = RGBColor(0xFF, 0xFF, 0xFF)   # White text
COLOUR_SECTION_BG  = RGBColor(0xEB, 0xF4, 0xFF)   # Light blue — section heading rows
COLOUR_YES_BG      = RGBColor(0xC6, 0xF6, 0xD5)   # Light green
COLOUR_PARTIAL_BG  = RGBColor(0xFE, 0xF3, 0xC7)   # Light amber
COLOUR_NO_BG       = RGBColor(0xFE, 0xD7, 0xD7)   # Light red
COLOUR_NA_BG       = RGBColor(0xF0, 0xF0, 0xF0)   # Light grey
COLOUR_HIGH        = RGBColor(0x22, 0x54, 0x3D)   # Dark green
COLOUR_LOW         = RGBColor(0x74, 0x2A, 0x2A)   # Dark red
COLOUR_BODY_TEXT   = RGBColor(0x2D, 0x37, 0x48)   # Dark slate

RATING_BG = {
    "Yes":     COLOUR_YES_BG,
    "Partial": COLOUR_PARTIAL_BG,
    "No":      COLOUR_NO_BG,
    "N/A":     COLOUR_NA_BG,
}

RATING_ICON = {
    "Yes":     "✅  Yes",
    "Partial": "⚠️  Partial",
    "No":      "❌  No",
    "N/A":     "—  N/A",
}


# ── Low-level XML helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_val = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_val)
    tcPr.append(shd)


def _set_cell_borders(cell, border_style: str = "single", size: int = 4) -> None:
    """Apply thin borders to all sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), border_style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_paragraph(cell, text: str, bold: bool = False,
                    font_size: int = 9, color: RGBColor | None = None,
                    align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Replace a cell's content with a single styled paragraph."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Tight vertical padding
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "20")
    spacing.set(qn("w:after"), "20")
    pPr.append(spacing)


# ── Document-level helpers ───────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _body(doc: Document, text: str, bold: bool = False,
          color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _labelled_para(doc: Document, label: str, value: str,
                   label_color: RGBColor | None = None,
                   value_color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(label + "  ")
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    if label_color:
        label_run.font.color.rgb = label_color
    val_run = p.add_run(value)
    val_run.font.size = Pt(10)
    if value_color:
        val_run.font.color.rgb = value_color


# ── Main export function ─────────────────────────────────────────────────────

def export_appraisal_to_docx(
    appraisals: List[PaperAppraisal],
    output_path: str = "quality_appraisal.docx",
) -> Path:
    """
    Write quality appraisal results to a formatted Word document.

    Each paper gets:
    - A header with reference, study type, and quality score badge
    - A 20-criterion table with colour-coded ratings and justifications
    - Strengths and limitations narratives

    Returns the Path to the saved file.
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Document title ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("REST Quality Assessment")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOUR_HEADER_BG

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(
        "Rapid Evidence Synthesis Team — 20-Criterion Quality Appraisal Tool"
    )
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_paragraph()  # spacer

    # ── One section per paper ─────────────────────────────────────────────────
    for paper_idx, appraisal in enumerate(appraisals, 1):
        # Paper header
        hdr_p = doc.add_paragraph()
        hdr_run = hdr_p.add_run(f"Paper {paper_idx}")
        hdr_run.font.size = Pt(14)
        hdr_run.font.bold = True
        hdr_run.font.color.rgb = COLOUR_HEADER_BG

        # Reference
        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(appraisal.article_reference)
        ref_run.font.size = Pt(9)
        ref_run.font.italic = True
        ref_run.font.color.rgb = COLOUR_BODY_TEXT

        doc.add_paragraph()  # spacer

        # Quality score badge row
        score_p = doc.add_paragraph()
        score_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        study_run = score_p.add_run(f"Study design: {appraisal.study_type}     ")
        study_run.font.size = Pt(10)
        study_run.font.bold = True
        study_run.font.color.rgb = COLOUR_BODY_TEXT

        score_run = score_p.add_run(f"Quality score: {appraisal.quality_score}")
        score_run.font.size = Pt(10)
        score_run.font.bold = True

        doc.add_paragraph()  # spacer

        # ── Criteria table ────────────────────────────────────────────────────
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # Column widths: #, Question, Rating, Justification
        col_widths = [Cm(1.0), Cm(7.5), Cm(2.5), Cm(8.5)]
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

        # Header row
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            _set_cell_bg(cell, COLOUR_HEADER_BG)
            _set_cell_borders(cell)
        _cell_paragraph(hdr_cells[0], "#",             bold=True, color=COLOUR_HEADER_FG, font_size=9)
        _cell_paragraph(hdr_cells[1], "Criterion",     bold=True, color=COLOUR_HEADER_FG, font_size=9)
        _cell_paragraph(hdr_cells[2], "Rating",        bold=True, color=COLOUR_HEADER_FG, font_size=9)
        _cell_paragraph(hdr_cells[3], "Justification", bold=True, color=COLOUR_HEADER_FG, font_size=9)

        # Data rows
        for criterion in appraisal.criteria:
            row = table.add_row()
            cells = row.cells

            rating_bg = RATING_BG.get(criterion.rating, COLOUR_NA_BG)
            for cell in cells:
                _set_cell_bg(cell, rating_bg)
                _set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            _cell_paragraph(
                cells[0],
                str(criterion.criterion_id),
                bold=True,
                font_size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _cell_paragraph(cells[1], criterion.question, font_size=9)
            _cell_paragraph(
                cells[2],
                RATING_ICON.get(criterion.rating, criterion.rating),
                bold=True,
                font_size=9,
            )
            _cell_paragraph(cells[3], criterion.justification, font_size=9)

        doc.add_paragraph()  # spacer after table

        # ── Strengths & Limitations ───────────────────────────────────────────
        _labelled_para(
            doc,
            "Strengths:",
            appraisal.strengths,
            label_color=COLOUR_HIGH,
        )
        _labelled_para(
            doc,
            "Limitations:",
            appraisal.limitations,
            label_color=COLOUR_LOW,
        )

        # Page break between papers (not after last one)
        if paper_idx < len(appraisals):
            doc.add_page_break()

    # ── Save ─────────────────────────────────────────────────────────────────
    out = Path(output_path)
    doc.save(out)
    print(f"\n✅ Quality appraisal saved to: {out.resolve()}")
    return out
